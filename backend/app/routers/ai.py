from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from app.schemas import SummarizeRequest, ReflectionOut, ProjectCreate
from app import models
from app.database import get_db
from app.deps import get_current_user
from typing import Optional, List
from pydantic import BaseModel
import os
import io
import json
import base64
from google import genai
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from PIL import Image
from docx import Document
from fpdf import FPDF

load_dotenv()

router = APIRouter()

# Initialize new Gemini SDK client
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
client = None
if GEMINI_API_KEY:
    client = genai.Client(api_key=GEMINI_API_KEY)
else:
    print("WARNING: GEMINI_API_KEY not found in .env")
class SummarizeWithSaveRequest(BaseModel):
    text: str
    project_title: Optional[str] = None
    save_to_project: bool = True

def get_gemini_summary(text: str) -> str:
    try:
        if not client:
            return "Gemini API key not configured."
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=f"Please provide a concise, highly professional summary of the following text, suitable for documenting a project or personal reflection. Focus on the main points and actionable insights if any:\n\n{text}"
        )
        return response.text
    except Exception as e:
        print(f"Gemini API Error: {e}")
        return text[:200] + "..." if len(text) > 200 else text

@router.post("/summarize")
async def summarize_and_save(
    request: SummarizeWithSaveRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Summarize text using Google Gemini API and optionally save to database.
    """
    summary_text = get_gemini_summary(request.text)
    
    if request.save_to_project:
        project_title = request.project_title or f"Project - {request.text[:30]}"
        db_project = models.Project(title=project_title, description="Auto-created via AI summarization", user_id=current_user.id)
        db.add(db_project)
        db.commit()
        db.refresh(db_project)
        
        db_reflection = models.Reflection(input_text=request.text, summary=summary_text, project_id=db_project.id)
        db.add(db_reflection)
        db.commit()
        db.refresh(db_reflection)
        
        return {
            "status": "success", "summary": summary_text, 
            "project_id": db_project.id, "project_title": db_project.title,
            "reflection_id": db_reflection.id, "saved_to_database": True
        }
    
    return {"status": "success", "summary": summary_text, "saved_to_database": False}


@router.post("/upload_and_analyze")
async def upload_and_analyze(
    files: List[UploadFile] = File(...),
    project_title: Optional[str] = Form(None),
    output_format: str = Form("default"),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Receives multiple images or PDFs, extracts text/visuals, uses Gemini to 
    generate a structured JSON response, and converts to requested format.
    """
    try:
        extracted_text = ""
        gemini_inputs = []
        file_names = []
        
        for file in files:
            contents = await file.read()
            file_ext = os.path.splitext(file.filename)[1].lower()
            file_names.append(file.filename)
            
            if file_ext == '.pdf':
                pdf_reader = PdfReader(io.BytesIO(contents))
                for page in pdf_reader.pages:
                    extracted_text += page.extract_text() + "\n"
            elif file_ext in ['.png', '.jpg', '.jpeg', '.webp']:
                image = Image.open(io.BytesIO(contents))
                gemini_inputs.append(image)
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported format for {file.filename}")
                
        if extracted_text:
            gemini_inputs.append(extracted_text)
            
        prompt = """
        You are an expert academic and software engineering assistant.
        Analyze the provided document(s) or image(s), representing a student's project, notes, or assignment.
        
        Generate a comprehensive, structured response containing exactly the following JSON structure:
        {
            "latex_report": "A complete, well-formatted LaTeX document (with \\documentclass{article}, title, author, \\begin{document}, sections, etc.) that beautifully summarizes and formats the project details.",
            "key_takeaways": ["Point 1", "Point 2", "Point 3"],
            "skills_developed": ["Skill 1", "Skill 2"],
            "suggested_next_steps": ["Step 1", "Step 2"],
            "difficulty_rating": 7
        }
        
        Ensure the output is strictly valid JSON. Do not include markdown codeblocks around the JSON.
        """
        gemini_inputs.append(prompt)
        
        if not client:
            raise HTTPException(status_code=500, detail="Gemini client not initialized")
            
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=gemini_inputs,
            config={"response_mime_type": "application/json"}
        )
        
        try:
            result_json = json.loads(response.text)
        except Exception as e:
            print("Failed to parse JSON from Gemini:", response.text)
            raise HTTPException(status_code=500, detail="Failed to parse AI response.")
            
        # Create Project and Reflection
        p_title = project_title or f"Analyzed Documents - {', '.join(file_names)[:30]}"
        db_project = models.Project(title=p_title, description="Uploaded Files Analysis via Gemini", user_id=current_user.id)
        db.add(db_project)
        db.commit()
        db.refresh(db_project)
        
        db_reflection = models.Reflection(
            input_text=", ".join(file_names), 
            summary=result_json.get("latex_report", "No LaTeX generated"), 
            project_id=db_project.id
        )
        db.add(db_reflection)
        db.commit()
        db.refresh(db_reflection)
        
        download_data = None
        
        if output_format == "docx":
            document = Document()
            document.add_heading(p_title, 0)
            document.add_heading('Key Takeaways', level=1)
            for tk in result_json.get("key_takeaways", []):
                document.add_paragraph(tk, style='List Bullet')
            document.add_heading('Skills Developed', level=1)
            for sk in result_json.get("skills_developed", []):
                document.add_paragraph(sk, style='List Bullet')
            document.add_heading('LaTeX Output', level=1)
            document.add_paragraph(result_json.get("latex_report", ""))
                
            doc_io = io.BytesIO()
            document.save(doc_io)
            doc_io.seek(0)
            download_data = base64.b64encode(doc_io.read()).decode("utf-8")
            
        elif output_format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt=p_title, ln=True, align='C')
            pdf.ln(10)
            pdf.cell(200, 10, txt="Key Takeaways:", ln=True)
            for tk in result_json.get("key_takeaways", []):
                pdf.multi_cell(0, 10, txt=f"- {tk}")
            pdf.ln(5)
            pdf.cell(200, 10, txt="Skills:", ln=True)
            for sk in result_json.get("skills_developed", []):
                pdf.cell(200, 10, txt=f"- {sk}", ln=True)
                
            pdf_bytes = pdf.output(dest='S')
            download_data = base64.b64encode(pdf_bytes).decode("utf-8")
        
        return {
            "status": "success",
            "project_id": db_project.id,
            "project_title": db_project.title,
            "analysis": result_json,
            "download_data": download_data,
            "format": output_format
        }
        
    except Exception as e:
        print(f"Error during upload and analyze: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
